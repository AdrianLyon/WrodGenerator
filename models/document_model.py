from docx import Document
import pandas as pd

class DocumentModel:
    def __init__(self):
        self.document = None
        self.file_path = None

    def load_document(self, file_path):
        self.document = Document(file_path)
        self.file_path = file_path

    def save_document(self, save_path):
        if self.document:
            self.document.save(save_path)

    def fill_form(self, document, data):
        paragraphs = document.paragraphs

        # Lógica de validación
        errors = self.validate_data(data)
        if errors:
            return False, errors

        # Extraer datos desde el diccionario 'data'
        first_name = data.get('name', '')
        last_name = data.get('LastNames', '')
        age = data.get('Age', '')
        day = data.get('Day', '')
        month = data.get('Month', '')
        year = data.get('Year', '')
        tutor = data.get('Tutor', '')

        # Modificar el documento con los datos
        for paragraph in paragraphs:
            if 'FirstNames:' in paragraph.text and 'LastNames:' in paragraph.text:
                paragraph.text = f"FirstNames: {first_name} LastNames: {last_name}"
            elif 'Age:' in paragraph.text and 'Day' in paragraph.text:
                paragraph.text = f"Age: {age} Day {day} Month {month} Year {year}"
            elif 'Tutor:' in paragraph.text:
                paragraph.text = f"Tutor: {tutor}"

        return True, None

    def validate_data(self, data):
        errors = []
        if not data.get('name'):
            errors.append("First Names cannot be empty.")
        if not data.get('LastNames'):
            errors.append("Last Names cannot be empty.")
        age = data.get('Age', "")
        if not age.isdigit() or int(age) <= 0:
            errors.append("Age must be a positive number.")
        day = data.get('Day', "")
        if not day.isdigit() or not (1 <= int(day) <= 31):
            errors.append("Day must be a number between 1 and 31.")
        month = data.get('Month', "")
        if not month.isdigit() or not (1 <= int(month) <= 12):
            errors.append("Month must be a number between 1 and 12.")
        year = data.get('Year', "")
        if not year.isdigit() or int(year) <= 0:
            errors.append("Year must be a positive number.")
        return errors

    def generate_documents_from_excel(self, excel_path):
        df = pd.read_excel(excel_path)
        generated_files = []
        for index, row in df.iterrows():
            document = Document(self.file_path)
            data = {
                'FirstNames': str(row['name']),
                'LastNames': str(row['LastNames']),
                'Age': str(row['Age']),
                'Day': str(row['Day']),
                'Month': str(row['Month']),
                'Year': str(row['Year']),
                'Tutor': str(row['Tutor'])
            }
            success, _ = self.fill_form(document, data)
            if success:
                generated_files.append(document)
        return generated_files
